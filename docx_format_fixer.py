#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Arabic Document Format Fixer
Fixes RTL, alignment, and formatting issues in Arabic .docx and .odt files
Supports both Microsoft Word (DOCX) and OpenDocument (ODF/ODT) formats
"""

import os
import sys
import zipfile
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from xml.etree import ElementTree as ET
import argparse
from encoding_fixer import ArabicEncodingFixer
from odf_handler import ODFHandler


class ArabicDocxFixer:
    """Handles fixing Arabic document formatting issues"""

    # Arabic fonts to use (in order of preference)
    ARABIC_FONTS = ['Arial', 'Traditional Arabic', 'Simplified Arabic']

    def __init__(self, folder_path, dry_run=False, fix_encoding=True):
        """Initialize the fixer with a folder path"""
        self.folder_path = Path(folder_path)
        self.fixed_count = 0
        self.error_count = 0
        self.fixes_applied = {}
        self.dry_run = dry_run
        self.fix_encoding = fix_encoding
        self.encoding_fixer = ArabicEncodingFixer() if fix_encoding else None
        self.odf_handler = ODFHandler(fix_encoding=fix_encoding)

    def is_arabic_text(self, text):
        """Check if text contains Arabic characters"""
        if not text:
            return False
        arabic_chars = sum(1 for char in text if '\u0600' <= char <= '\u06FF')
        return arabic_chars > len(text) * 0.3  # More than 30% Arabic characters

    def set_rtl_paragraph(self, paragraph):
        """Set paragraph to RTL (right-to-left) direction"""
        pPr = paragraph._element.get_or_add_pPr()

        # Check if bidi already exists
        existing_bidi = pPr.find(qn('w:bidi'))
        if existing_bidi is None:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

    def set_rtl_run(self, run):
        """Set run to RTL direction"""
        rPr = run._element.get_or_add_rPr()

        # Check if rtl already exists
        existing_rtl = rPr.find(qn('w:rtl'))
        if existing_rtl is None:
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), '1')
            rPr.append(rtl)

    def fix_text_encoding_in_paragraph(self, paragraph, doc_fixes):
        """Fix Mojibake encoding issues in paragraph text"""
        if not paragraph.runs:
            return

        encoding_fixed = False

        for run in paragraph.runs:
            if run.text:
                original_text = run.text
                fixed_text = self.encoding_fixer.clean_text(original_text)

                if fixed_text != original_text:
                    run.text = fixed_text
                    encoding_fixed = True

        if encoding_fixed:
            doc_fixes['encoding'] += 1

    def fix_numbered_headers(self, paragraph):
        """Fix manually numbered headers - handles multiple patterns"""
        import re

        text = paragraph.text.strip()

        # Pattern 1: "النطاق.2" → "2. النطاق" (NO SPACE before dot)
        pattern1 = r'^(.+?)\.(\d+(?:\.\d+)*)$'

        # Pattern 2: "المجال .2" → "2. المجال" (WITH SPACE before dot)
        pattern2 = r'^(.+?)\s+\.(\d+(?:\.\d+)*)$'

        match = re.match(pattern1, text) or re.match(pattern2, text)

        if match:
            arabic_text = match.group(1).strip()
            number = match.group(2).strip()

            # Check if text is actually Arabic (not just a file name or URL)
            if not self.is_arabic_text(arabic_text):
                return False

            # Reconstruct in proper RTL format: "2. النطاق"
            # Number first, then dot, then space, then Arabic text
            new_text = f"{number}. {arabic_text}"

            # Preserve formatting by getting the first run's properties
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                is_bold = first_run.font.bold
                is_italic = first_run.font.italic

                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''

                # Remove extra runs
                while len(paragraph.runs) > 1:
                    paragraph._element.remove(paragraph.runs[-1]._element)

                # Set new text with preserved formatting
                paragraph.runs[0].text = new_text
                paragraph.runs[0].font.name = font_name
                paragraph.runs[0].font.size = font_size
                paragraph.runs[0].font.bold = is_bold
                paragraph.runs[0].font.italic = is_italic

                return True

        return False

    def fix_paragraph_formatting(self, paragraph, doc_fixes):
        """Fix paragraph alignment and RTL direction - NO font changes"""
        if not paragraph.text.strip():
            return

        # Fix encoding issues FIRST (before any other processing)
        if self.fix_encoding and self.encoding_fixer:
            self.fix_text_encoding_in_paragraph(paragraph, doc_fixes)

        is_arabic = self.is_arabic_text(paragraph.text)

        # Try to fix numbered headers first (like "المجال .2")
        if is_arabic:
            if self.fix_numbered_headers(paragraph):
                doc_fixes['fonts'] += 1  # Reusing fonts counter for numbered headers

        # Apply RTL and alignment to ALL paragraphs (not just Arabic)
        # Set RTL direction
        self.set_rtl_paragraph(paragraph)
        doc_fixes['rtl_paragraphs'] += 1

        # Set right alignment
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc_fixes['alignments'] += 1

        # Fix numbered/bulleted lists - set list numbering to RTL
        pPr = paragraph._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            # This paragraph has list numbering - ensure bidi is set
            existing_bidi = pPr.find(qn('w:bidi'))
            if existing_bidi is None:
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)

        if is_arabic:
            # Only set RTL for runs, NO font changes
            for run in paragraph.runs:
                # Set RTL for the run
                self.set_rtl_run(run)

    def reverse_table_columns(self, table):
        """Reverse the order of columns in a table (mirror for RTL) by swapping cell XML elements"""
        for row in table.rows:
            # Get all cell elements
            cells = list(row.cells)
            num_cells = len(cells)

            if num_cells <= 1:
                continue  # Nothing to reverse

            # Store cell XML elements
            cell_elements = [cell._element for cell in cells]

            # Reverse the list
            cell_elements_reversed = list(reversed(cell_elements))

            # Get the parent row element
            row_element = row._element

            # Find the table cell elements in the row
            tc_elements = row_element.findall(qn('w:tc'))

            # Remove all cells from row
            for tc in tc_elements:
                row_element.remove(tc)

            # Add them back in reversed order
            for tc in cell_elements_reversed:
                row_element.append(tc)

    def fix_table_formatting(self, table, doc_fixes):
        """Fix table cell alignment and RTL direction for Arabic content"""
        # First, check if table has Arabic content
        has_arabic = False
        for row in table.rows:
            for cell in row.cells:
                if self.is_arabic_text(cell.text.strip()):
                    has_arabic = True
                    break
            if has_arabic:
                break

        # If table has Arabic content, reverse columns
        if has_arabic:
            self.reverse_table_columns(table)
            doc_fixes['table_cells'] += 1

        # Then fix formatting
        for row in table.rows:
            for cell in row.cells:
                # Fix each paragraph in the cell
                for paragraph in cell.paragraphs:
                    self.fix_paragraph_formatting(paragraph, doc_fixes)

    def fix_bullet_formatting(self, paragraph, doc_fixes):
        """Fix bullet point formatting for Arabic text"""
        if paragraph.style.name.startswith('List') or '•' in paragraph.text:
            if self.is_arabic_text(paragraph.text):
                # Ensure RTL and right alignment
                self.set_rtl_paragraph(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Fix spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                doc_fixes['bullets'] += 1

    def get_document_content_summary(self, doc):
        """Get a summary of document content for validation"""
        paragraph_count = len(doc.paragraphs)
        total_text_length = sum(len(p.text) for p in doc.paragraphs)
        table_count = len(doc.tables)
        table_cells_count = sum(len(row.cells) for table in doc.tables for row in table.rows)

        return {
            'paragraph_count': paragraph_count,
            'total_text_length': total_text_length,
            'table_count': table_count,
            'table_cells_count': table_cells_count
        }

    def align_right_via_xml(self, doc_path, output_path):
        """
        FINAL STEP: Apply right alignment to ALL content by directly editing XML.
        This is the "Select All + Align Right" equivalent.
        """
        print(f"  ✓ Applying final right alignment via XML...")
        
        temp_dir = doc_path.parent / f"_temp_align_{doc_path.stem}"
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
        temp_dir.mkdir()

        try:
            # Extract docx
            with zipfile.ZipFile(output_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # Fix document.xml
            document_xml_path = temp_dir / 'word' / 'document.xml'
            if document_xml_path.exists():
                with open(document_xml_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # Parse XML
                root = ET.fromstring(content)
                w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

                count = 0
                # Find ALL paragraphs and set alignment to right
                for para in root.iter(f'{w}p'):
                    pPr = para.find(f'{w}pPr')
                    if pPr is None:
                        pPr = ET.Element(f'{w}pPr')
                        para.insert(0, pPr)

                    # Remove existing alignment
                    jc = pPr.find(f'{w}jc')
                    if jc is not None:
                        pPr.remove(jc)
                    
                    # Add right alignment
                    jc = ET.SubElement(pPr, f'{w}jc')
                    jc.set(f'{w}val', 'right')
                    count += 1

                # Save
                with open(document_xml_path, 'w', encoding='utf-8') as f:
                    f.write(ET.tostring(root, encoding='unicode'))

                print(f"    → Aligned {count} paragraphs to the right")

            # Repack
            if output_path.exists():
                output_path.unlink()
                
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root_dir, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = Path(root_dir) / file
                        arcname = file_path.relative_to(temp_dir)
                        zipf.write(file_path, arcname)

            return count

        finally:
            if temp_dir.exists():
                shutil.rmtree(temp_dir)

    def fix_odf_document(self, doc_path):
        """Fix a single ODF document (.odt)"""
        try:
            output_path = doc_path.parent / f"{doc_path.stem}_fixed{doc_path.suffix}"

            if self.dry_run:
                print(f"\n[DRY RUN] Would process ODF: {doc_path.name}")
                return True

            # Use ODF handler
            doc_fixes = self.odf_handler.fix_document(doc_path, output_path)

            self.fixed_count += 1
            self.fixes_applied[doc_path.name] = doc_fixes

            return True

        except Exception as e:
            print(f"✗ Error processing ODF {doc_path.name}: {str(e)}")
            import traceback
            traceback.print_exc()
            self.error_count += 1
            return False

    def fix_document(self, doc_path):
        """Fix a single document and return statistics"""
        # Check file type and route to appropriate handler
        if doc_path.suffix.lower() == '.odt':
            return self.fix_odf_document(doc_path)

        # DOCX handling below
        try:
            print(f"\nProcessing: {doc_path.name}")
            doc = Document(doc_path)

            # Get original content summary for validation
            original_summary = self.get_document_content_summary(doc)
            print(f"  Original: {original_summary['paragraph_count']} paragraphs, "
                  f"{original_summary['total_text_length']} chars, "
                  f"{original_summary['table_count']} tables")

            # Track fixes for this document
            doc_fixes = {
                'rtl_paragraphs': 0,
                'alignments': 0,
                'fonts': 0,
                'table_cells': 0,
                'bullets': 0,
                'encoding': 0
            }

            # Fix paragraphs
            for paragraph in doc.paragraphs:
                self.fix_paragraph_formatting(paragraph, doc_fixes)
                self.fix_bullet_formatting(paragraph, doc_fixes)

            # Fix tables
            for table in doc.tables:
                self.fix_table_formatting(table, doc_fixes)

            # Validate content hasn't been lost
            fixed_summary = self.get_document_content_summary(doc)

            # Check if content was preserved
            content_preserved = (
                original_summary['paragraph_count'] == fixed_summary['paragraph_count'] and
                original_summary['total_text_length'] == fixed_summary['total_text_length'] and
                original_summary['table_count'] == fixed_summary['table_count']
            )

            if not content_preserved:
                print(f"  ⚠ WARNING: Content mismatch detected!")
                print(f"  After fix: {fixed_summary['paragraph_count']} paragraphs, "
                      f"{fixed_summary['total_text_length']} chars, "
                      f"{fixed_summary['table_count']} tables")
                print(f"  This document may need manual review.")

            # Save fixed document (unless dry-run)
            output_path = doc_path.parent / f"{doc_path.stem}_fixed{doc_path.suffix}"

            if self.dry_run:
                print(f"  [DRY RUN] Would save to: {output_path.name}")
            else:
                # First save with python-docx
                doc.save(output_path)
                
                # FINAL STEP: Apply right alignment to everything via XML
                self.align_right_via_xml(doc_path, output_path)
                
                print(f"✓ Fixed and saved to: {output_path.name}")

            print(f"  - RTL paragraphs: {doc_fixes['rtl_paragraphs']}")
            print(f"  - Alignments: {doc_fixes['alignments']}")
            print(f"  - Fonts: {doc_fixes['fonts']}")
            print(f"  - Table cells: {doc_fixes['table_cells']}")
            print(f"  - Bullets: {doc_fixes['bullets']}")
            print(f"  - Encoding fixes: {doc_fixes['encoding']}")

            if content_preserved:
                print(f"  ✓ Content validation: PASSED")
            else:
                print(f"  ✗ Content validation: FAILED")

            self.fixed_count += 1
            self.fixes_applied[doc_path.name] = doc_fixes

            return True

        except Exception as e:
            print(f"✗ Error processing {doc_path.name}: {str(e)}")
            import traceback
            traceback.print_exc()
            self.error_count += 1
            return False

    def process_folder(self, recursive=True):
        """Process all .docx and .odt files in the folder and optionally in subfolders"""
        if not self.folder_path.exists():
            print(f"Error: Folder not found: {self.folder_path}")
            return

        if not self.folder_path.is_dir():
            print(f"Error: Not a directory: {self.folder_path}")
            return

        # Find all .docx and .odt files (recursively if enabled)
        if recursive:
            # Search in current folder and all subfolders
            docx_files = [
                f for f in self.folder_path.rglob("*.docx")
                if not f.name.startswith("~$") and "_fixed" not in f.name
            ]
            odt_files = [
                f for f in self.folder_path.rglob("*.odt")
                if not f.name.startswith(".~") and "_fixed" not in f.name
            ]
            print(f"\nSearching recursively in: {self.folder_path}")
        else:
            # Search only in current folder
            docx_files = [
                f for f in self.folder_path.glob("*.docx")
                if not f.name.startswith("~$") and "_fixed" not in f.name
            ]
            odt_files = [
                f for f in self.folder_path.glob("*.odt")
                if not f.name.startswith(".~") and "_fixed" not in f.name
            ]
            print(f"\nSearching in: {self.folder_path}")

        # Combine both file types
        all_files = docx_files + odt_files

        if not all_files:
            print(f"No .docx or .odt files found")
            return

        print(f"Found {len(all_files)} document(s) to process:")
        print(f"  - {len(docx_files)} .docx file(s)")
        print(f"  - {len(odt_files)} .odt file(s)")
        print("=" * 60)

        # Process each document
        for doc_path in all_files:
            # Show relative path for better clarity
            try:
                rel_path = doc_path.relative_to(self.folder_path)
                if rel_path.parent != Path('.'):
                    print(f"\n[{rel_path.parent}]")
            except:
                pass

            self.fix_document(doc_path)

        # Print summary
        self.print_summary()

    def print_summary(self):
        """Print a summary of the fixes applied"""
        print("\n" + "=" * 60)
        print("SUMMARY REPORT")
        print("=" * 60)
        print(f"Documents processed: {self.fixed_count + self.error_count}")
        print(f"Successfully fixed: {self.fixed_count}")
        print(f"Errors: {self.error_count}")

        if self.fixes_applied:
            print("\nDetailed fixes:")
            for doc_name, fixes in self.fixes_applied.items():
                total_fixes = sum(fixes.values())
                print(f"\n  {doc_name}: {total_fixes} total fixes")
                for fix_type, count in fixes.items():
                    if count > 0:
                        print(f"    - {fix_type}: {count}")


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description='Fix Arabic formatting issues in Word and OpenDocument files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python docx_format_fixer.py /path/to/documents
  python docx_format_fixer.py "C:\\Users\\Documents\\Arabic Docs"

Supported formats:
  - Microsoft Word (.docx)
  - OpenDocument Text (.odt)
        """
    )

    parser.add_argument(
        'folder_path',
        nargs='?',
        help='Path to folder containing .docx and .odt files'
    )

    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Preview changes without saving files'
    )

    parser.add_argument(
        '--no-recursive',
        action='store_true',
        help='Do not search in subfolders (only process current folder)'
    )

    parser.add_argument(
        '--no-encoding-fix',
        action='store_true',
        help='Skip Mojibake encoding fixes (e.g., ÀB → آب)'
    )

    args = parser.parse_args()

    # Get folder path from argument or user input
    if args.folder_path:
        folder_path = args.folder_path
    else:
        print("Arabic Document Format Fixer (DOCX & ODF)")
        print("=" * 60)
        folder_path = input("Enter the folder path containing .docx/.odt files: ").strip()

        if not folder_path:
            print("Error: No folder path provided")
            sys.exit(1)

    # Remove quotes if present
    folder_path = folder_path.strip('"').strip("'")

    # Create fixer and process
    fix_encoding = not args.no_encoding_fix
    fixer = ArabicDocxFixer(folder_path, dry_run=args.dry_run, fix_encoding=fix_encoding)

    if args.dry_run:
        print("\n*** DRY RUN MODE - No files will be modified ***\n")

    if fix_encoding:
        print("✓ Mojibake encoding fix enabled (ÀB → آب)")
    else:
        print("⚠ Encoding fix disabled")

    # Process with or without recursion
    recursive = not args.no_recursive
    fixer.process_folder(recursive=recursive)


if __name__ == "__main__":
    main()