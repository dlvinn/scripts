#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Complete Align Right Tool (with RTL/Bidi Support) - Debug Version
"""

import sys
import argparse
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def align_document_right(doc_path, output_path=None):
    """Align all content in a document to the right (Text + Direction + Tables + Lists)"""
    print(f"\n{'='*60}")
    print(f"Processing: {doc_path}")
    print(f"File exists: {doc_path.exists()}")
    print(f"File is file: {doc_path.is_file()}")
    print(f"{'='*60}")

    try:
        # Load document
        doc = Document(str(doc_path))  # Convert to string for compatibility
        print(f"✓ Document loaded successfully")
    except Exception as e:
        print(f"✗ Error opening file: {e}")
        import traceback
        traceback.print_exc()
        return None

    count = 0
    list_count = 0
    table_count = 0

    # 1. Align all regular paragraphs
    print(f"\nProcessing {len(doc.paragraphs)} paragraphs...")
    for paragraph in doc.paragraphs:
        try:
            # Visual Alignment (Right side of page)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Text Direction (Right-to-Left logic)
            try:
                paragraph.paragraph_format.bidi = True
            except Exception as e:
                print(f"  Warning: Could not set bidi for paragraph: {e}")

            # Fix numbered/bulleted lists
            if paragraph.style.name.startswith('List') or paragraph._element.pPr is not None:
                try:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    pPr = paragraph._element.get_or_add_pPr()
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        existing_bidi = pPr.find(qn('w:bidi'))
                        if existing_bidi is None:
                            bidi = OxmlElement('w:bidi')
                            bidi.set(qn('w:val'), '1')
                            pPr.append(bidi)
                        list_count += 1
                except Exception as e:
                    print(f"  Warning: Could not fix list: {e}")

            count += 1
        except Exception as e:
            print(f"  Error processing paragraph: {e}")

    # 2. Align all tables and their contents
    print(f"\nProcessing {len(doc.tables)} tables...")
    for table in doc.tables:
        try:
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            table_count += 1
        except Exception as e:
            print(f"  Warning: Could not align table: {e}")

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    try:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        try:
                            paragraph.paragraph_format.bidi = True
                        except:
                            pass

                        # Fix numbered lists in tables
                        if paragraph.style.name.startswith('List') or paragraph._element.pPr is not None:
                            try:
                                from docx.oxml.ns import qn
                                from docx.oxml import OxmlElement

                                pPr = paragraph._element.get_or_add_pPr()
                                numPr = pPr.find(qn('w:numPr'))
                                if numPr is not None:
                                    existing_bidi = pPr.find(qn('w:bidi'))
                                    if existing_bidi is None:
                                        bidi = OxmlElement('w:bidi')
                                        bidi.set(qn('w:val'), '1')
                                        pPr.append(bidi)
                                    list_count += 1
                            except:
                                pass

                        count += 1
                    except Exception as e:
                        print(f"  Error processing table cell: {e}")

    # Save
    if output_path is None:
        output_path = doc_path.parent / f"{doc_path.stem}_aligned{doc_path.suffix}"
    else:
        output_path = Path(output_path)

    print(f"\nSaving to: {output_path}")
    try:
        doc.save(str(output_path))  # Convert to string for compatibility
        print(f"\n{'='*60}")
        print(f"SUCCESS!")
        print(f"{'='*60}")
        print(f"✓ Aligned {count} paragraphs to right (with RTL support)")
        print(f"✓ Processed {table_count} tables")
        if list_count > 0:
            print(f"✓ Fixed {list_count} numbered/bulleted lists")
        print(f"✓ Saved to: {output_path}")
        print(f"{'='*60}\n")
        return output_path
    except Exception as e:
        print(f"✗ Error saving file: {e}")
        import traceback
        traceback.print_exc()
        return None


def process_folder(folder_path, recursive=True):
    """Process all .docx files in a folder and subfolders"""
    folder = Path(folder_path).resolve()  # Get absolute path
    
    print(f"\n{'='*60}")
    print(f"Folder path: {folder}")
    print(f"Folder exists: {folder.exists()}")
    print(f"Is directory: {folder.is_dir()}")
    print(f"{'='*60}")

    if not folder.exists():
        print(f"✗ Error: Folder not found: {folder}")
        return

    if not folder.is_dir():
        print(f"✗ Error: Not a directory: {folder}")
        return

    # Find all .docx files
    if recursive:
        docx_files = [
            f for f in folder.rglob("*.docx")
            if not f.name.startswith("~$") and "_aligned" not in f.name
        ]
        print(f"\nSearching recursively in: {folder}")
    else:
        docx_files = [
            f for f in folder.glob("*.docx")
            if not f.name.startswith("~$") and "_aligned" not in f.name
        ]
        print(f"\nSearching in: {folder}")

    print(f"Found {len(docx_files)} document(s) to process")
    
    if not docx_files:
        print(f"\n✗ No .docx files found")
        print(f"Note: Files starting with '~$' or containing '_aligned' are skipped")
        return

    success_count = 0
    error_count = 0

    for doc_path in docx_files:
        result = align_document_right(doc_path)
        if result:
            success_count += 1
        else:
            error_count += 1

    # Summary
    print(f"\n{'='*60}")
    print(f"SUMMARY")
    print(f"{'='*60}")
    print(f"Successfully aligned: {success_count}")
    print(f"Errors: {error_count}")
    print(f"{'='*60}\n")


def process_single_file(file_path, output_path=None):
    """Process a single file"""
    file_path = Path(file_path).resolve()  # Get absolute path

    print(f"\n{'='*60}")
    print(f"File path: {file_path}")
    print(f"File exists: {file_path.exists()}")
    print(f"Is file: {file_path.is_file()}")
    print(f"Extension: {file_path.suffix}")
    print(f"{'='*60}")

    if not file_path.exists():
        print(f"✗ Error: File not found: {file_path}")
        return

    if not file_path.suffix.lower() == '.docx':
        print(f"✗ Error: File must be a .docx file (got: {file_path.suffix})")
        return

    if output_path:
        output_path = Path(output_path).resolve()
        
    align_document_right(file_path, output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Align all content in Word documents to the right (RTL Support)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python script.py /path/to/folder          # Process all .docx in folder
  python script.py /path/to/file.docx       # Process single file
  python script.py --file document.docx     # Process single file
  python script.py --file doc.docx --output out.docx  # With custom output
        """
    )

    parser.add_argument(
        'path',
        nargs='?',
        help='Path to folder or file'
    )

    parser.add_argument(
        '--file',
        help='Process a single file instead of a folder'
    )

    parser.add_argument(
        '--output',
        help='Output file path (only for single file mode)'
    )

    parser.add_argument(
        '--no-recursive',
        action='store_true',
        help='Don\'t search subfolders (folder mode only)'
    )

    args = parser.parse_args()

    print("\n" + "="*60)
    print("Align Right Tool (RTL/Arabic Support)")
    print("="*60)

    # Determine mode based on arguments
    if args.file:
        # Single file mode via flag
        file_path = args.file.strip('"').strip("'")
        output_path = args.output.strip('"').strip("'") if args.output else None
        process_single_file(file_path, output_path)
        
    elif args.path:
        # Path argument provided
        path = Path(args.path.strip('"').strip("'")).resolve()
        if path.is_file():
            process_single_file(path)
        else:
            recursive = not args.no_recursive
            process_folder(path, recursive=recursive)
            
    else:
        # No arguments provided -> Interactive Mode
        print("\n1. Process folder (all .docx files)")
        print("2. Process single file")
        choice = input("\nSelect option (1 or 2): ").strip()

        if choice == "1":
            folder_input = input("Enter folder path: ").strip().strip('"').strip("'")
            if folder_input:
                process_folder(folder_input)
        elif choice == "2":
            file_input = input("Enter file path: ").strip().strip('"').strip("'")
            output_input = input("Enter output path (optional, press Enter to skip): ").strip().strip('"').strip("'")
            if file_input:
                process_single_file(file_input, output_input if output_input else None)
        else:
            print("✗ Invalid choice")

if __name__ == "__main__":
    main()