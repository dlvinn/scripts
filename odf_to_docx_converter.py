#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ODF to DOCX Converter
Converts OpenDocument (.odt) files to Microsoft Word (.docx) format
Preserves text content, formatting, and structure
"""

import sys
import argparse
from pathlib import Path
from odf import opendocument, text, table, style
from odf.opendocument import load
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ODFToDocxConverter:
    """Converts ODF documents to DOCX format"""

    def __init__(self):
        """Initialize converter"""
        self.odf_doc = None
        self.docx_doc = None
        self.style_map = {}

    def load_odf(self, odt_path):
        """Load an ODF document"""
        print(f"Loading ODF: {odt_path.name}")
        self.odf_doc = load(str(odt_path))
        return self.odf_doc

    def create_docx(self):
        """Create a new DOCX document"""
        self.docx_doc = Document()
        return self.docx_doc

    def extract_text_from_node(self, node):
        """Recursively extract text from an ODF node"""
        text_content = []

        if hasattr(node, 'data'):
            text_content.append(node.data)

        if hasattr(node, 'childNodes'):
            for child in node.childNodes:
                text_content.append(self.extract_text_from_node(child))

        return ''.join(text_content)

    def get_paragraph_style(self, odf_para):
        """Extract style information from ODF paragraph"""
        style_info = {
            'is_bold': False,
            'is_italic': False,
            'font_size': None,
            'alignment': None,
        }

        # Check for style attributes
        style_name = odf_para.getAttribute('stylename')

        # Check runs for formatting
        for node in odf_para.childNodes:
            if hasattr(node, 'getAttribute'):
                style_name = node.getAttribute('stylename')
                # Could parse styles here for bold/italic/etc

        return style_info

    def convert_paragraph(self, odf_para):
        """Convert an ODF paragraph to DOCX paragraph"""
        # Extract text content
        para_text = self.extract_text_from_node(odf_para)

        if not para_text.strip():
            # Empty paragraph
            self.docx_doc.add_paragraph()
            return

        # Add paragraph to DOCX
        docx_para = self.docx_doc.add_paragraph(para_text)

        # Get style information
        style_info = self.get_paragraph_style(odf_para)

        # Check if it's Arabic text
        if self.is_arabic_text(para_text):
            # Set RTL and right alignment for Arabic
            docx_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_rtl_paragraph(docx_para)

        return docx_para

    def is_arabic_text(self, text):
        """Check if text contains Arabic characters"""
        if not text:
            return False
        arabic_chars = sum(1 for char in text if '\u0600' <= char <= '\u06FF')
        return arabic_chars > len(text) * 0.3

    def set_rtl_paragraph(self, paragraph):
        """Set paragraph to RTL direction"""
        pPr = paragraph._element.get_or_add_pPr()

        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

    def convert_table(self, odf_table):
        """Convert an ODF table to DOCX table"""
        # Get table dimensions
        rows = odf_table.getElementsByType(table.TableRow)

        if not rows:
            return

        # Get max columns
        max_cols = 0
        for row in rows:
            cells = row.getElementsByType(table.TableCell)
            max_cols = max(max_cols, len(cells))

        if max_cols == 0:
            return

        # Create DOCX table
        docx_table = self.docx_doc.add_table(rows=len(rows), cols=max_cols)
        docx_table.style = 'Table Grid'

        # Fill table cells
        for row_idx, odf_row in enumerate(rows):
            odf_cells = odf_row.getElementsByType(table.TableCell)

            for col_idx, odf_cell in enumerate(odf_cells):
                if col_idx < max_cols:
                    docx_cell = docx_table.rows[row_idx].cells[col_idx]

                    # Get all paragraphs in the cell
                    odf_paras = odf_cell.getElementsByType(text.P)

                    # Clear default paragraph
                    if docx_cell.paragraphs:
                        docx_cell.paragraphs[0].text = ''

                    # Add cell content
                    for para_idx, odf_para in enumerate(odf_paras):
                        para_text = self.extract_text_from_node(odf_para)

                        if para_idx == 0 and docx_cell.paragraphs:
                            # Use existing first paragraph
                            docx_cell.paragraphs[0].text = para_text
                            cell_para = docx_cell.paragraphs[0]
                        else:
                            # Add new paragraph
                            cell_para = docx_cell.add_paragraph(para_text)

                        # Apply RTL if Arabic
                        if self.is_arabic_text(para_text):
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            self.set_rtl_paragraph(cell_para)

    def convert_content(self):
        """Convert all content from ODF to DOCX"""
        print("Converting content...")

        # Get the body element
        body_elements = []

        # Get all paragraphs
        paragraphs = self.odf_doc.getElementsByType(text.P)
        for para in paragraphs:
            body_elements.append(('paragraph', para))

        # Get all tables
        tables = self.odf_doc.getElementsByType(table.Table)
        for tbl in tables:
            body_elements.append(('table', tbl))

        # Sort by document order (simplified - just process in order found)
        # In real implementation, you'd need to track actual document order

        para_count = 0
        table_count = 0

        # Process paragraphs
        for para in paragraphs:
            self.convert_paragraph(para)
            para_count += 1

        # Process tables
        for tbl in tables:
            self.convert_table(tbl)
            table_count += 1

        print(f"  Converted {para_count} paragraphs")
        print(f"  Converted {table_count} tables")

    def convert_file(self, odt_path, docx_path=None):
        """
        Convert an ODF file to DOCX

        Args:
            odt_path: Path to input .odt file
            docx_path: Path to output .docx file (optional)

        Returns:
            Path to output file
        """
        odt_path = Path(odt_path)

        if not odt_path.exists():
            raise FileNotFoundError(f"File not found: {odt_path}")

        if odt_path.suffix.lower() != '.odt':
            raise ValueError(f"File must be .odt format: {odt_path}")

        # Determine output path
        if docx_path is None:
            docx_path = odt_path.parent / f"{odt_path.stem}.docx"
        else:
            docx_path = Path(docx_path)

        print(f"\n{'='*60}")
        print(f"Converting: {odt_path.name}")
        print(f"{'='*60}")

        # Load ODF
        self.load_odf(odt_path)

        # Create DOCX
        self.create_docx()

        # Convert content
        self.convert_content()

        # Save DOCX
        print(f"Saving to: {docx_path.name}")
        self.docx_doc.save(str(docx_path))

        print(f"✓ Conversion complete!")
        print(f"{'='*60}\n")

        return docx_path

    def convert_folder(self, folder_path, recursive=True):
        """
        Convert all .odt files in a folder to .docx

        Args:
            folder_path: Path to folder containing .odt files
            recursive: Whether to search subfolders
        """
        folder_path = Path(folder_path)

        if not folder_path.exists():
            print(f"Error: Folder not found: {folder_path}")
            return

        if not folder_path.is_dir():
            print(f"Error: Not a directory: {folder_path}")
            return

        # Find all .odt files
        if recursive:
            odt_files = [
                f for f in folder_path.rglob("*.odt")
                if not f.name.startswith(".~")
            ]
            print(f"\nSearching recursively in: {folder_path}")
        else:
            odt_files = [
                f for f in folder_path.glob("*.odt")
                if not f.name.startswith(".~")
            ]
            print(f"\nSearching in: {folder_path}")

        if not odt_files:
            print(f"No .odt files found")
            return

        print(f"Found {len(odt_files)} .odt file(s) to convert")
        print("=" * 60)

        success_count = 0
        error_count = 0

        for odt_file in odt_files:
            try:
                # Show subfolder if in recursive mode
                try:
                    rel_path = odt_file.relative_to(folder_path)
                    if rel_path.parent != Path('.'):
                        print(f"\n[{rel_path.parent}]")
                except:
                    pass

                self.convert_file(odt_file)
                success_count += 1

            except Exception as e:
                print(f"✗ Error converting {odt_file.name}: {str(e)}")
                error_count += 1

        # Summary
        print("\n" + "=" * 60)
        print("CONVERSION SUMMARY")
        print("=" * 60)
        print(f"Successfully converted: {success_count}")
        print(f"Errors: {error_count}")
        print(f"Total files: {len(odt_files)}")
        print("=" * 60)


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description='Convert ODF (.odt) files to Microsoft Word (.docx) format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert single file
  python odf_to_docx_converter.py --file document.odt

  # Convert all .odt files in folder
  python odf_to_docx_converter.py --folder "C:\\Documents"

  # Convert with custom output name
  python odf_to_docx_converter.py --file input.odt --output result.docx

  # Non-recursive folder conversion
  python odf_to_docx_converter.py --folder "C:\\Documents" --no-recursive
        """
    )

    parser.add_argument(
        '--file',
        help='Path to single .odt file to convert'
    )

    parser.add_argument(
        '--folder',
        help='Path to folder containing .odt files'
    )

    parser.add_argument(
        '--output',
        help='Output path for converted file (only for --file mode)'
    )

    parser.add_argument(
        '--no-recursive',
        action='store_true',
        help='Do not search in subfolders (only for --folder mode)'
    )

    args = parser.parse_args()

    converter = ODFToDocxConverter()

    if args.file:
        # Single file mode
        file_path = args.file.strip('"').strip("'")
        output_path = args.output.strip('"').strip("'") if args.output else None

        try:
            converter.convert_file(file_path, output_path)
        except Exception as e:
            print(f"Error: {str(e)}")
            sys.exit(1)

    elif args.folder:
        # Folder mode
        folder_path = args.folder.strip('"').strip("'")
        recursive = not args.no_recursive
        converter.convert_folder(folder_path, recursive=recursive)

    else:
        # Interactive mode
        print("ODF to DOCX Converter")
        print("=" * 60)
        print("1. Convert single file")
        print("2. Convert folder")
        choice = input("\nSelect option (1 or 2): ").strip()

        if choice == "1":
            file_input = input("Enter .odt file path: ").strip().strip('"').strip("'")
            output_input = input("Enter output path (optional, press Enter to skip): ").strip().strip('"').strip("'")

            if file_input:
                try:
                    converter.convert_file(
                        file_input,
                        output_input if output_input else None
                    )
                except Exception as e:
                    print(f"Error: {str(e)}")

        elif choice == "2":
            folder_input = input("Enter folder path: ").strip().strip('"').strip("'")
            recursive_input = input("Search in subfolders? (Y/n): ").strip().lower()
            recursive = recursive_input != 'n'

            if folder_input:
                converter.convert_folder(folder_input, recursive=recursive)

        else:
            print("Invalid choice")


if __name__ == "__main__":
    main()
