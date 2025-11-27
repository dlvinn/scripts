#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ODF to DOCX Converter - Single File Version
Converts a single .odt file to .docx format
"""

import sys
import io
from pathlib import Path
from odf import opendocument, text, table
from odf.opendocument import load
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Ensure UTF-8 output
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')


def is_arabic_text(text):
    """Check if text contains Arabic characters"""
    if not text:
        return False
    arabic_chars = sum(1 for char in text if '\u0600' <= char <= '\u06FF')
    return arabic_chars > len(text) * 0.3


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)


def extract_text_from_node(node):
    """Recursively extract text from an ODF node"""
    text_content = []
    if hasattr(node, 'data'):
        text_content.append(node.data)
    if hasattr(node, 'childNodes'):
        for child in node.childNodes:
            text_content.append(extract_text_from_node(child))
    return ''.join(text_content)


def convert_paragraph(odf_para, docx_doc):
    """Convert an ODF paragraph to DOCX"""
    para_text = extract_text_from_node(odf_para)

    if not para_text.strip():
        docx_doc.add_paragraph()
        return

    docx_para = docx_doc.add_paragraph(para_text)

    if is_arabic_text(para_text):
        docx_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(docx_para)


def convert_table(odf_table, docx_doc):
    """Convert an ODF table to DOCX"""
    rows = odf_table.getElementsByType(table.TableRow)
    if not rows:
        return

    max_cols = 0
    for row in rows:
        cells = row.getElementsByType(table.TableCell)
        max_cols = max(max_cols, len(cells))

    if max_cols == 0:
        return

    docx_table = docx_doc.add_table(rows=len(rows), cols=max_cols)
    docx_table.style = 'Table Grid'

    for row_idx, odf_row in enumerate(rows):
        odf_cells = odf_row.getElementsByType(table.TableCell)

        for col_idx, odf_cell in enumerate(odf_cells):
            if col_idx < max_cols:
                docx_cell = docx_table.rows[row_idx].cells[col_idx]
                odf_paras = odf_cell.getElementsByType(text.P)

                if docx_cell.paragraphs:
                    docx_cell.paragraphs[0].text = ''

                for para_idx, odf_para in enumerate(odf_paras):
                    para_text = extract_text_from_node(odf_para)

                    if para_idx == 0 and docx_cell.paragraphs:
                        docx_cell.paragraphs[0].text = para_text
                        cell_para = docx_cell.paragraphs[0]
                    else:
                        cell_para = docx_cell.add_paragraph(para_text)

                    if is_arabic_text(para_text):
                        cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        set_rtl_paragraph(cell_para)


def convert_file(odt_path, docx_path=None):
    """Convert a single ODT file to DOCX"""
    odt_path = Path(odt_path)

    if not odt_path.exists():
        print(f"✗ Error: File not found: {odt_path}")
        return False

    if odt_path.suffix.lower() != '.odt':
        print(f"✗ Error: File must be .odt format")
        return False

    if docx_path is None:
        docx_path = odt_path.parent / f"{odt_path.stem}.docx"
    else:
        docx_path = Path(docx_path)

    print(f"\n{'='*60}")
    print(f"Converting: {odt_path.name}")
    print(f"{'='*60}")

    try:
        # Load ODF
        print(f"✓ Loading ODF document...")
        odf_doc = load(str(odt_path))

        # Create DOCX
        print(f"✓ Creating DOCX document...")
        docx_doc = Document()

        # Get content
        paragraphs = odf_doc.getElementsByType(text.P)
        tables = odf_doc.getElementsByType(table.Table)

        print(f"✓ Found {len(paragraphs)} paragraphs")
        print(f"✓ Found {len(tables)} tables")

        # Convert paragraphs
        print(f"✓ Converting paragraphs...")
        for para in paragraphs:
            convert_paragraph(para, docx_doc)

        # Convert tables
        if tables:
            print(f"✓ Converting tables...")
            for tbl in tables:
                convert_table(tbl, docx_doc)

        # Save
        print(f"✓ Saving to: {docx_path.name}")
        docx_doc.save(str(docx_path))

        print(f"\n{'='*60}")
        print(f"✓ SUCCESS!")
        print(f"{'='*60}")
        print(f"Input:  {odt_path}")
        print(f"Output: {docx_path}")
        print(f"{'='*60}\n")

        return True

    except Exception as e:
        print(f"\n✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main entry point"""
    import argparse

    parser = argparse.ArgumentParser(
        description='Convert a single ODF (.odt) file to DOCX format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert with auto-generated output name
  python odf_to_docx_single.py document.odt

  # Convert with custom output name
  python odf_to_docx_single.py input.odt --output result.docx

  # Full paths
  python odf_to_docx_single.py "C:\\Docs\\report.odt" --output "C:\\Docs\\report.docx"
        """
    )

    parser.add_argument(
        'file',
        help='Path to .odt file to convert'
    )

    parser.add_argument(
        '--output',
        help='Output .docx path (optional, defaults to same name)'
    )

    args = parser.parse_args()

    # Convert
    file_path = args.file.strip('"').strip("'")
    output_path = args.output.strip('"').strip("'") if args.output else None

    success = convert_file(file_path, output_path)

    if not success:
        sys.exit(1)


if __name__ == "__main__":
    main()
