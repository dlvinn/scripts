#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Mojibake Discovery Tool
Scans .docx files to find all corrupted encoding characters
"""

import sys
from pathlib import Path
from docx import Document
from encoding_fixer import ArabicEncodingFixer


def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file (paragraphs and tables)"""
    doc = Document(docx_path)
    texts = []

    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)

    return texts


def scan_folder(folder_path, recursive=True):
    """
    Scan all .docx files in a folder for Mojibake characters

    Args:
        folder_path (Path): Path to folder containing .docx files
        recursive (bool): Whether to search in subfolders

    Returns:
        dict: Statistics and findings
    """
    folder_path = Path(folder_path)

    if not folder_path.exists():
        print(f"Error: Folder not found: {folder_path}")
        return None

    # Find all .docx files
    if recursive:
        docx_files = [
            f for f in folder_path.rglob("*.docx")
            if not f.name.startswith("~$") and "_fixed" not in f.name
        ]
    else:
        docx_files = [
            f for f in folder_path.glob("*.docx")
            if not f.name.startswith("~$") and "_fixed" not in f.name
        ]

    if not docx_files:
        print(f"No .docx files found in: {folder_path}")
        return None

    print(f"\n{'='*70}")
    print(f"SCANNING {len(docx_files)} DOCUMENT(S) FOR MOJIBAKE")
    print(f"{'='*70}\n")

    # Collect all text samples
    all_texts = []
    file_samples = {}  # Store sample corrupted text per file

    fixer = ArabicEncodingFixer()

    for doc_path in docx_files:
        try:
            print(f"Scanning: {doc_path.name}... ", end='', flush=True)
            texts = extract_text_from_docx(doc_path)

            # Find texts with Mojibake
            corrupted_samples = []
            for text in texts:
                discovery = fixer.discover_mojibake_chars([text])
                if discovery['found_mojibake'] or discovery['unmapped_chars']:
                    corrupted_samples.append(text[:100])  # First 100 chars

            if corrupted_samples:
                file_samples[doc_path.name] = corrupted_samples[:3]  # Keep first 3 samples
                print(f"✗ Found {len(corrupted_samples)} corrupted text(s)")
            else:
                print(f"✓ Clean")

            all_texts.extend(texts)

        except Exception as e:
            print(f"Error: {str(e)}")

    # Generate comprehensive report
    print(f"\n{'='*70}")
    print("OVERALL DISCOVERY REPORT")
    print(f"{'='*70}")

    fixer.print_discovery_report(all_texts)

    # Show sample corrupted texts per file
    if file_samples:
        print(f"\n{'='*70}")
        print("SAMPLE CORRUPTED TEXTS BY FILE")
        print(f"{'='*70}\n")

        for filename, samples in file_samples.items():
            print(f"📄 {filename}")
            for i, sample in enumerate(samples, 1):
                # Show first 80 chars
                display_text = sample[:80] + ('...' if len(sample) > 80 else '')
                print(f"   {i}. {display_text}")

                # Show fixed version
                fixed = fixer.clean_text(sample[:80])
                print(f"      → {fixed}")
            print()

    return {
        'total_files': len(docx_files),
        'files_with_issues': len(file_samples),
        'all_texts': all_texts
    }


def main():
    """Main entry point"""
    print("Mojibake Discovery Tool for Arabic Documents")
    print("=" * 70)

    if len(sys.argv) > 1:
        folder_path = sys.argv[1].strip('"').strip("'")
    else:
        folder_path = input("Enter folder path to scan: ").strip('"').strip("'")

    if not folder_path:
        print("Error: No folder path provided")
        sys.exit(1)

    # Ask about recursive search
    recursive_input = input("Search in subfolders too? (Y/n): ").strip().lower()
    recursive = recursive_input != 'n'

    # Scan
    results = scan_folder(folder_path, recursive=recursive)

    if results:
        print(f"\n{'='*70}")
        print("SCAN SUMMARY")
        print(f"{'='*70}")
        print(f"Total files scanned: {results['total_files']}")
        print(f"Files with encoding issues: {results['files_with_issues']}")
        print(f"Clean files: {results['total_files'] - results['files_with_issues']}")
        print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
