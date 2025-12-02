
import os
import argparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import nsdecls

def set_table_borders(table):
    """Adds borders to a table."""
    tbl_props = table._tbl.tblPr
    if tbl_props is None:
        tbl_props = OxmlElement('w:tblPr')
        table._tbl.insert(0, tbl_props)
    
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        tbl_borders.append(border)
    tbl_props.append(tbl_borders)

def copy_paragraph(paragraph, doc):
    """Copies a paragraph and its formatting."""
    new_para = doc.add_paragraph()
    for run in paragraph.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = 'Calibri'
        new_run.font.size = Pt(11)
        new_run.font.color.rgb = RGBColor(0, 0, 0)
    return new_para

def copy_table(table, doc):
    """Copies a table and its content."""
    new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns), style=table.style)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            # Clear the default paragraph in the new cell
            if len(new_cell.paragraphs) > 0:
                p = new_cell.paragraphs[0]
                p.clear()

            for para in cell.paragraphs:
                new_para = new_cell.add_paragraph()
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = 'Calibri'
                    new_run.font.size = Pt(11)
                    new_run.font.color.rgb = RGBColor(0, 0, 0)
    set_table_borders(new_table)
    return new_table


def process_document(template_path, source_path, output_path, new_name):
    """
    Creates a new document from a template, copies content from a source,
    and applies formatting.
    """
    print(f"Processing {source_path}...")
    
    template_doc = Document(template_path)
    source_doc = Document(source_path)

    # Clear existing paragraph and table content in the template's body
    elements_to_remove = []
    for el in template_doc.element.body:
        if el.tag.endswith('p') or el.tag.endswith('tbl'):
            elements_to_remove.append(el)
    for el in elements_to_remove:
        template_doc.element.body.remove(el)

    # Copy content from source to template
    for element in source_doc.element.body:
        if element.tag.endswith('p'):
            p = Paragraph(element, source_doc)
            copy_paragraph(p, template_doc)
        elif element.tag.endswith('tbl'):
            table = Table(element, source_doc)
            copy_table(table, template_doc)

    # Modify the footer
    for section in template_doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            # It's possible a paragraph has multiple runs
            # We need to combine them to check, but replace in the specific run
            full_text = ''.join(r.text for r in paragraph.runs)
            if 'PRO-015' in full_text:
                for run in paragraph.runs:
                    if 'PRO-015' in run.text:
                        run.text = run.text.replace('PRO-015', new_name)

    print(f"  Saving to {output_path}")
    template_doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Process DOCX files based on a template.")
    parser.add_argument("input_dir", help="The directory containing DOCX files to process.")
    args = parser.parse_args()

    template_file = os.path.join("files", "AAH-IMS-PRO-015_Calibration.docx")
    
    if not os.path.exists(template_file):
        print(f"Error: Template file not found at {template_file}")
        return

    for root, _, files in os.walk(args.input_dir):
        for file in files:
            if file.endswith(".docx"):
                source_file_path = os.path.join(root, file)
                
                output_filename = os.path.basename(source_file_path)
                new_name_for_footer = output_filename.replace("_fixed.docx", "").replace(".docx", "")

                if output_filename.endswith("_fixed.docx"):
                    output_filename = output_filename.replace("_fixed.docx", ".docx")
                
                output_file_path = os.path.join(root, output_filename)

                process_document(template_file, source_file_path, output_file_path, new_name_for_footer)

if __name__ == "__main__":
    main()
