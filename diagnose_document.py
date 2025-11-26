#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Diagnostic Tool
Analyzes Word documents to identify potential issues before fixing
"""

import sys
from pathlib import Path
from docx import Document
import argparse


def analyze_document(doc_path):
    """Analyze a document and print detailed information"""
    print(f"\n{'='*70}")
    print(f"DOCUMENT ANALYSIS: {doc_path.name}")
    print(f"{'='*70}\n")

    try:
        doc = Document(doc_path)

        # Overall statistics
        print(f"📊 OVERALL STATISTICS")
        print(f"{'─'*70}")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")

        # Paragraph analysis
        print(f"\n📝 PARAGRAPH ANALYSIS")
        print(f"{'─'*70}")

        non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        print(f"Non-empty paragraphs: {len(non_empty_paragraphs)}")

        # Count Arabic vs non-Arabic
        arabic_count = 0
        english_count = 0
        mixed_count = 0

        for p in non_empty_paragraphs:
            text = p.text.strip()
            arabic_chars = sum(1 for char in text if '\u0600' <= char <= '\u06FF')
            english_chars = sum(1 for char in text if char.isalpha() and char.isascii())

            if arabic_chars > english_chars * 2:
                arabic_count += 1
            elif english_chars > arabic_chars * 2:
                english_count += 1
            else:
                mixed_count += 1

        print(f"  - Arabic paragraphs: {arabic_count}")
        print(f"  - English paragraphs: {english_count}")
        print(f"  - Mixed paragraphs: {mixed_count}")

        # Show first 10 paragraphs with line numbers
        print(f"\n📄 PARAGRAPH CONTENT (First 10)")
        print(f"{'─'*70}")
        for idx, p in enumerate(doc.paragraphs[:10], 1):
            text = p.text.strip()
            if text:
                # Truncate long text
                display_text = text[:80] + "..." if len(text) > 80 else text
                print(f"{idx:3}. {display_text}")

        # Table analysis
        if doc.tables:
            print(f"\n📊 TABLE ANALYSIS")
            print(f"{'─'*70}")

            for table_idx, table in enumerate(doc.tables, 1):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                print(f"\nTable {table_idx}:")
                print(f"  - Rows: {rows}")
                print(f"  - Columns: {cols}")

                # Show first few cells
                if rows > 0 and cols > 0:
                    print(f"  - First row content:")
                    for cell_idx, cell in enumerate(table.rows[0].cells, 1):
                        cell_text = cell.text.strip()
                        if cell_text:
                            display_text = cell_text[:40] + "..." if len(cell_text) > 40 else cell_text
                            print(f"    Cell {cell_idx}: {display_text}")

        # Character statistics
        print(f"\n🔤 CHARACTER STATISTICS")
        print(f"{'─'*70}")

        all_text = '\n'.join(p.text for p in doc.paragraphs)
        total_chars = len(all_text)
        arabic_chars = sum(1 for char in all_text if '\u0600' <= char <= '\u06FF')
        english_chars = sum(1 for char in all_text if char.isalpha() and char.isascii())
        digit_chars = sum(1 for char in all_text if char.isdigit())
        space_chars = sum(1 for char in all_text if char.isspace())

        print(f"Total characters: {total_chars}")
        print(f"  - Arabic characters: {arabic_chars} ({arabic_chars/total_chars*100:.1f}%)")
        print(f"  - English characters: {english_chars} ({english_chars/total_chars*100:.1f}%)")
        print(f"  - Digits: {digit_chars} ({digit_chars/total_chars*100:.1f}%)")
        print(f"  - Spaces: {space_chars} ({space_chars/total_chars*100:.1f}%)")

        # Style analysis
        print(f"\n🎨 STYLE ANALYSIS")
        print(f"{'─'*70}")

        styles_used = {}
        for p in doc.paragraphs:
            style = p.style.name
            styles_used[style] = styles_used.get(style, 0) + 1

        print(f"Styles used: {len(styles_used)}")
        for style, count in sorted(styles_used.items(), key=lambda x: x[1], reverse=True)[:10]:
            print(f"  - {style}: {count}")

        # Formatting issues detection
        print(f"\n⚠️  POTENTIAL ISSUES")
        print(f"{'─'*70}")

        issues_found = False

        # Check for bullet points
        bullet_paragraphs = [p for p in doc.paragraphs if '•' in p.text]
        if bullet_paragraphs:
            print(f"✓ Found {len(bullet_paragraphs)} paragraphs with bullet points")
            issues_found = True

        # Check for mixed direction text
        mixed_direction = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                arabic_chars = sum(1 for char in text if '\u0600' <= char <= '\u06FF')
                english_chars = sum(1 for char in text if char.isalpha() and char.isascii())
                if arabic_chars > 5 and english_chars > 5:
                    mixed_direction.append(p)

        if mixed_direction:
            print(f"✓ Found {len(mixed_direction)} paragraphs with mixed Arabic/English")
            issues_found = True

        # Check for dates/numbers that might cause issues
        date_like = [p for p in doc.paragraphs if any(d in p.text for d in ['2025', '2024', '/', '-'])]
        if date_like:
            print(f"✓ Found {len(date_like)} paragraphs containing dates/numbers")
            issues_found = True

        if not issues_found:
            print("No obvious issues detected")

        print(f"\n{'='*70}\n")

    except Exception as e:
        print(f"❌ Error analyzing document: {str(e)}")
        import traceback
        traceback.print_exc()


def main():
    parser = argparse.ArgumentParser(
        description='Diagnose issues in Word documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python diagnose_document.py document.docx
  python diagnose_document.py "path/to/document.docx"
        """
    )

    parser.add_argument(
        'doc_path',
        nargs='?',
        help='Path to .docx file to analyze'
    )

    args = parser.parse_args()

    # Get file path
    if args.doc_path:
        doc_path = args.doc_path
    else:
        print("Document Diagnostic Tool")
        print("="*70)
        doc_path = input("Enter the path to the .docx file: ").strip()

        if not doc_path:
            print("Error: No file path provided")
            sys.exit(1)

    # Remove quotes
    doc_path = doc_path.strip('"').strip("'")
    doc_path = Path(doc_path)

    if not doc_path.exists():
        print(f"Error: File not found: {doc_path}")
        sys.exit(1)

    if not doc_path.suffix.lower() == '.docx':
        print(f"Error: File must be a .docx file")
        sys.exit(1)

    # Analyze
    analyze_document(doc_path)


if __name__ == "__main__":
    main()
