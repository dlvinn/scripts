#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Arabic Word Document Generator
Creates professional Arabic documents with proper RTL formatting
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import argparse
from datetime import datetime


class ArabicDocumentGenerator:
    """Generates professional Arabic Word documents"""

    # Default Arabic font
    ARABIC_FONT = 'Arial'

    def __init__(self, logo_path=None, company_name=None, content=None, output_path=None):
        """Initialize the generator with document parameters"""
        self.logo_path = Path(logo_path) if logo_path else None
        self.company_name = company_name or ""
        self.content = content or ""
        self.output_path = Path(output_path) if output_path else None
        self.doc = Document()

        # Set default page margins (narrower for professional look)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def set_rtl_paragraph(self, paragraph):
        """Set paragraph to RTL (right-to-left) direction"""
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def set_rtl_run(self, run):
        """Set run to RTL direction"""
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)

    def apply_arabic_formatting(self, paragraph, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
        """Apply proper Arabic formatting to a paragraph"""
        # Set RTL direction
        self.set_rtl_paragraph(paragraph)

        # Set alignment
        paragraph.alignment = alignment

        # Format runs
        for run in paragraph.runs:
            run.font.name = self.ARABIC_FONT
            run.font.size = Pt(font_size)
            run.font.bold = bold

            # Set RTL for run
            self.set_rtl_run(run)

            # Apply font to complex scripts (Arabic)
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), self.ARABIC_FONT)
            rFonts.set(qn('w:hAnsi'), self.ARABIC_FONT)
            rFonts.set(qn('w:cs'), self.ARABIC_FONT)
            rPr.append(rFonts)

    def add_header_with_logo(self):
        """Add logo and company name to the document header"""
        if not self.logo_path:
            print("Warning: No logo path provided, skipping logo")
            return

        if not self.logo_path.exists():
            print(f"Warning: Logo file not found at {self.logo_path}")
            return

        try:
            # Add logo - centered at top
            logo_paragraph = self.doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Calculate appropriate logo size
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(str(self.logo_path), width=Inches(1.5))

            # Add spacing after logo
            logo_paragraph.paragraph_format.space_after = Pt(6)

        except Exception as e:
            print(f"Error adding logo: {str(e)}")

    def add_company_name(self):
        """Add styled company name"""
        if not self.company_name:
            return

        # Company name - large, bold, centered
        company_paragraph = self.doc.add_paragraph()
        run = company_paragraph.add_run(self.company_name)

        # Apply styling
        self.apply_arabic_formatting(company_paragraph, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Add spacing
        company_paragraph.paragraph_format.space_after = Pt(12)
        company_paragraph.paragraph_format.space_before = Pt(6)

    def add_separator_line(self):
        """Add a horizontal separator line"""
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(12)

        # Add a border bottom to create a line
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_content(self):
        """Add main document content with proper formatting"""
        if not self.content:
            return

        # Split content by lines
        lines = self.content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                # Add empty paragraph for spacing
                self.doc.add_paragraph()
                continue

            # Check if it's a bullet point
            is_bullet = line.startswith('•') or line.startswith('-')

            # Add paragraph
            paragraph = self.doc.add_paragraph()

            # For bullets, use a style
            if is_bullet:
                # Remove the bullet character if present
                line = line.lstrip('•-').strip()
                # Add bullet
                paragraph.add_run('• ' + line)
                paragraph.paragraph_format.left_indent = Inches(0.25)
            else:
                paragraph.add_run(line)

            # Apply Arabic formatting
            self.apply_arabic_formatting(paragraph, font_size=12)

            # Set line spacing
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)

    def add_signature_section(self):
        """Add professional signature/approval section with three columns"""
        # Add spacing before signature section
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Create table with 3 columns
        table = self.doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Set column widths
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = Inches(1.8)

        # Headers (first row)
        headers = ['الإعداد', 'المراجعة', 'الاعتماد']
        header_row = table.rows[0]

        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.clear()
            run = cell_paragraph.add_run(header_text)
            run.font.bold = True

            # Apply Arabic formatting
            self.apply_arabic_formatting(cell_paragraph, font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Add shading to header
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E7E6E6')
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Row 2: الاسم (Name)
        labels_row = table.rows[1]
        for idx in range(3):
            cell = labels_row.cells[idx]
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.clear()
            cell_paragraph.add_run('الاسم: _________________')
            self.apply_arabic_formatting(cell_paragraph, font_size=11)

        # Row 3: التاريخ (Date)
        date_row = table.rows[2]
        for idx in range(3):
            cell = date_row.cells[idx]
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.clear()
            cell_paragraph.add_run('التاريخ: _________________')
            self.apply_arabic_formatting(cell_paragraph, font_size=11)

        # Row 4: التوقيع (Signature)
        signature_row = table.rows[3]
        for idx in range(3):
            cell = signature_row.cells[idx]
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.clear()
            cell_paragraph.add_run('التوقيع: _________________')
            self.apply_arabic_formatting(cell_paragraph, font_size=11)

        # Add spacing to table
        table_paragraph = table.rows[0].cells[0].paragraphs[0]
        table_paragraph.paragraph_format.space_before = Pt(12)

    def generate(self):
        """Generate the complete document"""
        print("Generating Arabic document...")

        # Add header with logo
        self.add_header_with_logo()

        # Add company name
        self.add_company_name()

        # Add separator
        self.add_separator_line()

        # Add main content
        self.add_content()

        # Add signature section
        self.add_signature_section()

        # Save document
        if not self.output_path:
            self.output_path = Path('arabic_document.docx')

        # Ensure directory exists
        self.output_path.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(self.output_path))
        print(f"✓ Document generated successfully: {self.output_path}")

        return self.output_path


def interactive_mode():
    """Run in interactive mode to collect inputs from user"""
    print("=" * 60)
    print("Arabic Document Generator - Interactive Mode")
    print("=" * 60)

    # Get logo path
    logo_path = input("\nEnter logo image path (or press Enter to skip): ").strip().strip('"').strip("'")
    if logo_path and not Path(logo_path).exists():
        print(f"Warning: Logo file not found at {logo_path}")
        logo_path = None

    # Get company name
    company_name = input("\nEnter company name (in Arabic): ").strip()

    # Get content
    print("\nEnter document content (in Arabic).")
    print("You can paste multiple lines. Type 'END' on a new line when done:")
    content_lines = []
    while True:
        line = input()
        if line.strip().upper() == 'END':
            break
        content_lines.append(line)
    content = '\n'.join(content_lines)

    # Get output path
    default_output = "arabic_document.docx"
    output_path = input(f"\nEnter output file path (default: {default_output}): ").strip().strip('"').strip("'")
    if not output_path:
        output_path = default_output

    return logo_path, company_name, content, output_path


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description='Generate professional Arabic Word documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Interactive mode
  python docx_generator.py

  # Command line mode
  python docx_generator.py --logo logo.png --company "شركة المثال" --content "المحتوى هنا" --output document.docx

  # With content from file
  python docx_generator.py --logo logo.png --company "شركة المثال" --content-file content.txt --output document.docx
        """
    )

    parser.add_argument('--logo', help='Path to company logo image')
    parser.add_argument('--company', help='Company name in Arabic')
    parser.add_argument('--content', help='Document content in Arabic')
    parser.add_argument('--content-file', help='Path to file containing document content')
    parser.add_argument('--output', help='Output file path (default: arabic_document.docx)')

    args = parser.parse_args()

    # Check if any arguments were provided
    if not any([args.logo, args.company, args.content, args.content_file, args.output]):
        # Run in interactive mode
        logo_path, company_name, content, output_path = interactive_mode()
    else:
        # Use command line arguments
        logo_path = args.logo
        company_name = args.company

        # Get content from file or argument
        if args.content_file:
            try:
                with open(args.content_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            except Exception as e:
                print(f"Error reading content file: {str(e)}")
                sys.exit(1)
        else:
            content = args.content

        output_path = args.output or 'arabic_document.docx'

    # Validate inputs
    if not company_name and not content:
        print("Error: At least company name or content must be provided")
        sys.exit(1)

    # Generate document
    try:
        generator = ArabicDocumentGenerator(
            logo_path=logo_path,
            company_name=company_name,
            content=content,
            output_path=output_path
        )
        generator.generate()

    except Exception as e:
        print(f"Error generating document: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
